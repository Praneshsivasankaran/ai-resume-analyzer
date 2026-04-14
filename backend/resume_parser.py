import os
import fitz  # PyMuPDF for PDFs
from docx import Document


def extract_text(file_path):
    """Extract text from PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path):
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text


def _extract_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        # Paragraphs
        text += "\n".join(p.text for p in doc.paragraphs)
        # Tables (skills sections often live here)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text